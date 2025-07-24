"""
Export functionality for multiple output formats.
"""

import os
import json
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any, List, Optional
import logging
from datetime import datetime

try:
    from ..utils.dependencies import safe_import, get_available_export_formats
except ImportError:
    from utils.dependencies import safe_import, get_available_export_formats

# Import optional dependencies using the centralized system
markdown, MARKDOWN_AVAILABLE = safe_import("markdown", "markdown")

# ReportLab imports
SimpleDocTemplate, REPORTLAB_AVAILABLE = safe_import("reportlab.platypus.SimpleDocTemplate", "reportlab")
letter = None
Paragraph = None
Spacer = None
getSampleStyleSheet = None
ParagraphStyle = None
inch = None

if REPORTLAB_AVAILABLE:
    letter, _ = safe_import("reportlab.lib.pagesizes.letter", "reportlab")
    Paragraph, _ = safe_import("reportlab.platypus.Paragraph", "reportlab")
    Spacer, _ = safe_import("reportlab.platypus.Spacer", "reportlab")
    getSampleStyleSheet, _ = safe_import("reportlab.lib.styles.getSampleStyleSheet", "reportlab")
    ParagraphStyle, _ = safe_import("reportlab.lib.styles.ParagraphStyle", "reportlab")
    inch, _ = safe_import("reportlab.lib.units.inch", "reportlab")

# Python-docx imports
Document, DOCX_AVAILABLE = safe_import("docx.Document", "python-docx")
if DOCX_AVAILABLE:
    Inches, _ = safe_import("docx.shared.Inches", "python-docx")


class ExporterBase(ABC):
    """Abstract base class for content exporters."""
    
    def __init__(self):
        """Initialize the exporter."""
        self.logger = logging.getLogger(__name__)
    
    @abstractmethod
    def export(self, content: str, output_path: Path, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """Export content to the specified format.
        
        Args:
            content: Content to export
            output_path: Path for the output file
            metadata: Optional metadata for the export
            
        Returns:
            True if export successful, False otherwise
        """
        pass
    
    @abstractmethod
    def get_file_extension(self) -> str:
        """Get the file extension for this format."""
        pass
    
    @abstractmethod
    def is_available(self) -> bool:
        """Check if the exporter is available (dependencies installed)."""
        pass


class MarkdownExporter(ExporterBase):
    """Export content to Markdown format with enhanced formatting."""
    
    def export(self, content: str, output_path: Path, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """Export content to Markdown format.
        
        Args:
            content: Content to export
            output_path: Path for the output file
            metadata: Optional metadata
            
        Returns:
            True if successful
        """
        try:
            # Ensure path has correct extension
            if not output_path.suffix:
                output_path = output_path.with_suffix('.md')
            elif output_path.suffix.lower() != '.md':
                output_path = output_path.with_suffix('.md')
            
            # Create markdown content with metadata
            markdown_content = self._format_markdown_content(content, metadata)
            
            # Write to file
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            self.logger.info(f"Exported content to Markdown: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to export to Markdown: {e}")
            return False
    
    def _format_markdown_content(self, content: str, metadata: Optional[Dict[str, Any]] = None) -> str:
        """Format content with Markdown styling and metadata.
        
        Args:
            content: Raw content
            metadata: Optional metadata
            
        Returns:
            Formatted Markdown content
        """
        formatted_content = []
        
        # Add title as main heading if available in metadata
        if metadata and "title" in metadata:
            formatted_content.append(f"# {metadata['title']}")
            formatted_content.append("")
        elif not metadata or "title" not in metadata:
            # Default title when no title in metadata
            formatted_content.append("# YouTube Transcript Export")
            formatted_content.append("")
        
        # Add other metadata as a section
        if metadata:
            formatted_content.append("---")
            formatted_content.append("# Document Metadata")
            formatted_content.append("")
            
            for key, value in metadata.items():
                if key == "title":
                    continue  # Already added as main heading
                elif key == "source_url":
                    formatted_content.append(f"**Source URL:** {value}")
                elif key == "generated_at":
                    formatted_content.append(f"**Generated:** {value}")
                elif key == "total_videos":
                    formatted_content.append(f"**Total Videos:** {value}")
                elif key == "processing_style":
                    formatted_content.append(f"**Processing Style:** {value}")
                else:
                    formatted_content.append(f"**{key.replace('_', ' ').title()}:** {value}")
            
            formatted_content.append("")
            formatted_content.append("---")
            formatted_content.append("")
        
        # Process content for better Markdown formatting
        lines = content.split('\n')
        in_transcript_section = False
        video_count = 0
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines at the beginning
            if not line and not formatted_content:
                continue
            
            # Detect video URLs and make them proper headers
            if line.startswith("Video URL:"):
                url = line.replace("Video URL:", "").strip()
                video_count += 1
                formatted_content.append("")
                formatted_content.append(f"## Video {video_count}")
                formatted_content.append("")
                formatted_content.append(f"🎥 [{url}]({url})")
                formatted_content.append("")
                in_transcript_section = True
                continue
            
            # Format section headers
            if line.lower().startswith(("summary:", "key points:", "main topics:")):
                formatted_content.append("")
                formatted_content.append(f"### {line}")
                formatted_content.append("")
                continue
            
            # Format bullet points
            if line.startswith("- ") or line.startswith("* "):
                formatted_content.append(line)
                continue
            
            # Regular content
            if line:
                formatted_content.append(line)
            else:
                formatted_content.append("")
        
        # Add table of contents if multiple videos
        if metadata and metadata.get("total_videos", 0) > 1:
            toc = self._generate_table_of_contents(formatted_content)
            # Insert TOC after metadata but before content
            metadata_end = 0
            for i, line in enumerate(formatted_content):
                if line == "---" and i > 0:
                    metadata_end = i + 2
                    break
            
            formatted_content[metadata_end:metadata_end] = toc
        
        return '\n'.join(formatted_content)
    
    def _generate_table_of_contents(self, content_lines: List[str]) -> List[str]:
        """Generate a table of contents from headers.
        
        Args:
            content_lines: List of content lines
            
        Returns:
            List of TOC lines
        """
        toc_lines = ["## 📋 Table of Contents", ""]
        
        for line in content_lines:
            if line.startswith("## ") and not line.startswith("## 📋"):
                # Extract header text and create anchor link
                header_text = line[3:].strip()  # Remove "## "
                anchor = header_text.lower().replace(" ", "-").replace("[", "").replace("]", "").replace("(", "").replace(")", "")
                toc_lines.append(f"- [{header_text}](#{anchor})")
        
        toc_lines.extend(["", "---", ""])
        return toc_lines
    
    def get_file_extension(self) -> str:
        """Get file extension."""
        return ".md"
    
    def is_available(self) -> bool:
        """Check if Markdown export is available."""
        return True  # Markdown export doesn't require external dependencies


class PDFExporter(ExporterBase):
    """Export content to PDF format with professional styling."""
    
    def export(self, content: str, output_path: Path, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """Export content to PDF format.
        
        Args:
            content: Content to export
            output_path: Path for output file
            metadata: Optional metadata
            
        Returns:
            True if successful
        """
        if not self.is_available():
            self.logger.error("PDF export not available - reportlab not installed")
            return False
        
        try:
            # Ensure correct extension
            if not output_path.suffix:
                output_path = output_path.with_suffix('.pdf')
            elif output_path.suffix.lower() != '.pdf':
                output_path = output_path.with_suffix('.pdf')
            
            # Check that all required imports are available
            if not SimpleDocTemplate:
                raise ImportError("ReportLab SimpleDocTemplate not available")
            
            # Create PDF
            output_path.parent.mkdir(parents=True, exist_ok=True)
            # Use letter if available, otherwise use a default size
            page_size = letter if letter else (8.5*72, 11*72)  # Default letter size in points
            doc = SimpleDocTemplate(str(output_path), pagesize=page_size)
            
            # Build content
            story = self._build_pdf_story(content, metadata)
            
            # Build PDF
            doc.build(story)
            
            self.logger.info(f"Exported content to PDF: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to export to PDF: {e}")
            return False
    
    def _build_pdf_story(self, content: str, metadata: Optional[Dict[str, Any]] = None) -> List[Any]:
        """Build the PDF story (content structure).
        
        Args:
            content: Raw content
            metadata: Optional metadata
            
        Returns:
            List of PDF elements
        """
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import Color
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            textColor=Color(0, 0, 0.5)  # dark blue
        )
        
        video_header_style = ParagraphStyle(
            'VideoHeader',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=Color(0, 0.5, 0)  # dark green
        )
        
        # Title
        story.append(Paragraph("YouTube Transcript Export", title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        if metadata:
            story.append(Paragraph("Document Information", styles['Heading3']))
            for key, value in metadata.items():
                formatted_key = key.replace('_', ' ').title()
                story.append(Paragraph(f"<b>{formatted_key}:</b> {value}", styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
        
        # Process content
        lines = content.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            
            if line.startswith("Video URL:"):
                # Finish current paragraph
                if current_paragraph:
                    story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                    current_paragraph = []
                
                # Add video header
                url = line.replace("Video URL:", "").strip()
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph(f"Video: {url}", video_header_style))
                story.append(Spacer(1, 0.1*inch))
                
            elif line:
                current_paragraph.append(line)
            else:
                # Empty line - finish current paragraph
                if current_paragraph:
                    story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
                    current_paragraph = []
        
        # Finish final paragraph
        if current_paragraph:
            story.append(Paragraph(' '.join(current_paragraph), styles['Normal']))
        
        return story
    
    def get_file_extension(self) -> str:
        """Get file extension."""
        return ".pdf"
    
    def is_available(self) -> bool:
        """Check if PDF export is available."""
        return REPORTLAB_AVAILABLE


class DocxExporter(ExporterBase):
    """Export content to Word document format."""
    
    def export(self, content: str, output_path: Path, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """Export content to DOCX format.
        
        Args:
            content: Content to export
            output_path: Path for output file
            metadata: Optional metadata
            
        Returns:
            True if successful
        """
        if not self.is_available():
            self.logger.error("DOCX export not available - python-docx not installed")
            return False
        
        try:
            # Ensure correct extension
            if not output_path.suffix:
                output_path = output_path.with_suffix('.docx')
            elif output_path.suffix.lower() != '.docx':
                output_path = output_path.with_suffix('.docx')
            
            # Create document
            doc = Document()
            
            # Add content
            self._build_docx_content(doc, content, metadata)
            
            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            self.logger.info(f"Exported content to DOCX: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to export to DOCX: {e}")
            return False
    
    def _build_docx_content(self, doc: Any, content: str, metadata: Optional[Dict[str, Any]] = None) -> None:
        """Build the DOCX document content.
        
        Args:
            doc: Document object
            content: Raw content
            metadata: Optional metadata
        """
        from docx.shared import Inches
        
        # Title
        title = doc.add_heading('YouTube Transcript Export', 0)
        
        # Metadata
        if metadata:
            doc.add_heading('Document Information', level=2)
            for key, value in metadata.items():
                formatted_key = key.replace('_', ' ').title()
                p = doc.add_paragraph()
                p.add_run(f"{formatted_key}: ").bold = True
                p.add_run(str(value))
        
        # Process content
        lines = content.split('\n')
        current_paragraph_lines = []
        
        for line in lines:
            line = line.strip()
            
            if line.startswith("Video URL:"):
                # Finish current paragraph
                if current_paragraph_lines:
                    doc.add_paragraph(' '.join(current_paragraph_lines))
                    current_paragraph_lines = []
                
                # Add video header
                url = line.replace("Video URL:", "").strip()
                doc.add_heading(f"Video: {url}", level=2)
                
            elif line:
                current_paragraph_lines.append(line)
            else:
                # Empty line - finish current paragraph
                if current_paragraph_lines:
                    doc.add_paragraph(' '.join(current_paragraph_lines))
                    current_paragraph_lines = []
        
        # Finish final paragraph
        if current_paragraph_lines:
            doc.add_paragraph(' '.join(current_paragraph_lines))
    
    def get_file_extension(self) -> str:
        """Get file extension."""
        return ".docx"
    
    def is_available(self) -> bool:
        """Check if DOCX export is available."""
        return DOCX_AVAILABLE


class HTMLExporter(ExporterBase):
    """Export content to HTML format with modern styling."""
    
    def export(self, content: str, output_path: Path, metadata: Optional[Dict[str, Any]] = None) -> bool:
        """Export content to HTML format.
        
        Args:
            content: Content to export
            output_path: Path for output file
            metadata: Optional metadata
            
        Returns:
            True if successful
        """
        try:
            # Ensure correct extension
            if not output_path.suffix:
                output_path = output_path.with_suffix('.html')
            elif output_path.suffix.lower() != '.html':
                output_path = output_path.with_suffix('.html')
            
            # Generate HTML
            html_content = self._generate_html_content(content, metadata)
            
            # Write to file
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.logger.info(f"Exported content to HTML: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to export to HTML: {e}")
            return False
    
    def _generate_html_content(self, content: str, metadata: Optional[Dict[str, Any]] = None) -> str:
        """Generate HTML content with styling.
        
        Args:
            content: Raw content
            metadata: Optional metadata
            
        Returns:
            Complete HTML document
        """
        html_parts = []
        
        # HTML header with modern styling
        html_parts.append('''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>YouTube Transcript Export</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9f9f9;
        }
        .container {
            background: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        h1 {
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
        }
        h2 {
            color: #27ae60;
            margin-top: 30px;
        }
        h3 {
            color: #8e44ad;
        }
        .metadata {
            background: #ecf0f1;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
        }
        .metadata strong {
            color: #2c3e50;
        }
        .video-section {
            border-left: 4px solid #3498db;
            padding-left: 20px;
            margin: 20px 0;
        }
        .video-url {
            background: #3498db;
            color: white;
            padding: 10px;
            border-radius: 5px;
            margin: 10px 0;
        }
        .video-url a {
            color: white;
            text-decoration: none;
        }
        .video-url a:hover {
            text-decoration: underline;
        }
        p {
            margin: 15px 0;
        }
        .footer {
            text-align: center;
            margin-top: 40px;
            padding: 20px;
            background: #34495e;
            color: white;
            border-radius: 5px;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>📺 YouTube Transcript Export</h1>
''')
        
        # Metadata section
        if metadata:
            html_parts.append('        <div class="metadata">')
            html_parts.append('            <h3>📋 Document Information</h3>')
            for key, value in metadata.items():
                formatted_key = key.replace('_', ' ').title()
                if key == "source_url":
                    html_parts.append(f'            <p><strong>{formatted_key}:</strong> <a href="{value}" target="_blank">{value}</a></p>')
                else:
                    html_parts.append(f'            <p><strong>{formatted_key}:</strong> {value}</p>')
            html_parts.append('        </div>')
        
        # Process content
        lines = content.split('\n')
        current_section = []
        in_video_section = False
        
        for line in lines:
            line = line.strip()
            
            if line.startswith("Video URL:"):
                # Finish current section
                if current_section and in_video_section:
                    html_parts.append('            <div class="video-content">')
                    for content_line in current_section:
                        if content_line:
                            html_parts.append(f'                <p>{self._escape_html(content_line)}</p>')
                    html_parts.append('            </div>')
                    html_parts.append('        </div>')
                    current_section = []
                
                # Start new video section
                url = line.replace("Video URL:", "").strip()
                html_parts.append('        <div class="video-section">')
                html_parts.append(f'            <div class="video-url">🎥 <a href="{url}" target="_blank">{url}</a></div>')
                in_video_section = True
                
            elif line:
                current_section.append(line)
            elif current_section:
                # Empty line but we have content - add it as a paragraph break
                current_section.append("")
        
        # Finish final section
        if current_section and in_video_section:
            html_parts.append('            <div class="video-content">')
            for content_line in current_section:
                if content_line:
                    html_parts.append(f'                <p>{self._escape_html(content_line)}</p>')
            html_parts.append('            </div>')
            html_parts.append('        </div>')
        
        # Footer
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        html_parts.append(f'''        <div class="footer">
            <p>Generated by YouTube Transcript Extractor on {current_time}</p>
        </div>
    </div>
</body>
</html>''')
        
        return '\n'.join(html_parts)
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters.
        
        Args:
            text: Text to escape
            
        Returns:
            HTML-escaped text
        """
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&#x27;'))
    
    def get_file_extension(self) -> str:
        """Get file extension."""
        return ".html"
    
    def is_available(self) -> bool:
        """Check if HTML export is available."""
        return True  # HTML export doesn't require external dependencies


class ExportManager:
    """Manager for handling multiple export formats."""
    
    def __init__(self):
        """Initialize export manager."""
        self.logger = logging.getLogger(__name__)
        self.exporters = {
            'markdown': MarkdownExporter(),
            'pdf': PDFExporter(),
            'docx': DocxExporter(),
            'html': HTMLExporter()
        }
    
    def get_available_formats(self) -> List[str]:
        """Get list of available export formats.
        
        Returns:
            List of format names
        """
        return [name for name, exporter in self.exporters.items() if exporter.is_available()]
    
    def export_content(
        self, 
        content: str, 
        format_name: str, 
        output_path: Path,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bool:
        """Export content to specified format.
        
        Args:
            content: Content to export
            format_name: Export format ('markdown', 'pdf', 'docx', 'html')
            output_path: Output file path
            metadata: Optional metadata
            
        Returns:
            True if successful
        """
        if format_name not in self.exporters:
            self.logger.error(f"Unknown export format: {format_name}")
            return False
        
        exporter = self.exporters[format_name]
        
        if not exporter.is_available():
            self.logger.error(f"Export format {format_name} not available (missing dependencies)")
            return False
        
        return exporter.export(content, output_path, metadata)
    
    def export_to_multiple_formats(
        self, 
        content: str, 
        base_output_path: Path,
        formats: List[str],
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, bool]:
        """Export content to multiple formats.
        
        Args:
            content: Content to export
            base_output_path: Base path for output files (without extension)
            formats: List of format names
            metadata: Optional metadata
            
        Returns:
            Dictionary mapping format names to success status
        """
        results = {}
        
        for format_name in formats:
            if format_name in self.exporters:
                exporter = self.exporters[format_name]
                extension = exporter.get_file_extension()
                output_path = base_output_path.with_suffix(extension)
                
                results[format_name] = self.export_content(content, format_name, output_path, metadata)
            else:
                results[format_name] = False
                self.logger.error(f"Unknown format: {format_name}")
        
        return results
    
    def get_missing_dependencies(self) -> Dict[str, List[str]]:
        """Get missing dependencies for export formats.
        
        Returns:
            Dictionary mapping format names to missing packages
        """
        missing = {}
        
        if not MARKDOWN_AVAILABLE:
            missing['markdown'] = ['markdown']
        
        if not REPORTLAB_AVAILABLE:
            missing['pdf'] = ['reportlab']
        
        if not DOCX_AVAILABLE:
            missing['docx'] = ['python-docx']
        
        return missing
